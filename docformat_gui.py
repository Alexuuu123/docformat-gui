#!/usr/bin/env python3
"""
公文格式处理工具 - 纸质感极简风格 v2
优化：更大图标、更好排版、卡片式选择
"""

import os
import sys
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path

# 添加scripts目录到路径
SCRIPT_DIR = Path(__file__).parent / "scripts"
sys.path.insert(0, str(SCRIPT_DIR))

from scripts.analyzer import analyze_punctuation, analyze_numbering, analyze_paragraph_format, analyze_font
from scripts.punctuation import process_document as fix_punctuation
from scripts.formatter import format_document, PRESETS


# ===== 设计系统 =====
class Theme:
    # 纸质色调
    BG = '#FBF9F6'              # 温暖米白纸张
    CARD = '#FFFFFF'            # 纯白卡片
    CARD_ALT = '#F7F4EF'        # 米黄卡片（推荐区）
    INPUT_BG = '#F2EFE9'        # 输入框背景（稍深米色）
    
    # 陶土红
    PRIMARY = '#BC4B26'         # 朱砂/印泥色
    PRIMARY_HOVER = '#A3421F'   # 悬停加深
    PRIMARY_LIGHT = '#F9F0EC'   # 极淡红
    
    # 文字
    TEXT = '#2E2E2E'            # 深炭灰
    TEXT_SECONDARY = '#6B6B6B'  # 次要文字
    TEXT_MUTED = '#A0A0A0'      # 禁用/占位
    
    # 边框与分隔
    BORDER = '#E8E4DE'          # 温暖灰边框
    BORDER_LIGHT = '#F0EDE8'    # 更浅边框
    BORDER_SELECTED = '#BC4B26' # 选中边框
    
    # 日志区
    LOG_BG = '#1A1A1A'
    LOG_TEXT = '#C8C8C8'
    LOG_SUCCESS = '#7CB87C'
    LOG_WARNING = '#D4A656'
    LOG_ERROR = '#CF6B6B'
    
    # 字体 - 宋体优先
    FONT_SERIF = ('Noto Serif SC', 'Source Han Serif SC', 'SimSun', 'PMingLiU', 'serif')
    
    # 间距
    SPACE_XS = 4
    SPACE_SM = 8
    SPACE_MD = 16
    SPACE_LG = 24
    SPACE_XL = 40


def get_font(size=12, weight='normal'):
    """获取宋体字体"""
    return (Theme.FONT_SERIF[0], size, weight)


# ===== 大尺寸线条图标 =====
class Icons:
    """用 Canvas 绘制的线条图标 - 48px 大尺寸"""
    
    @staticmethod
    def draw_magic(canvas, x, y, size=48, color='#2E2E2E'):
        """智能处理 - 魔法棒"""
        s = size
        lw = 2.5  # 线宽
        # 魔法棒主体
        canvas.create_line(x+s*0.15, y+s*0.85, x+s*0.65, y+s*0.35, fill=color, width=lw, capstyle='round')
        # 星星点缀
        stars = [(0.7, 0.2), (0.85, 0.35), (0.75, 0.5), (0.55, 0.15)]
        for px, py in stars:
            r = 3
            canvas.create_oval(x+s*px-r, y+s*py-r, x+s*px+r, y+s*py+r, fill=color, outline='')
        # 光芒线
        canvas.create_line(x+s*0.7, y+s*0.08, x+s*0.7, y+s*0.22, fill=color, width=1.5)
        canvas.create_line(x+s*0.9, y+s*0.28, x+s*0.78, y+s*0.35, fill=color, width=1.5)
    
    @staticmethod
    def draw_search(canvas, x, y, size=48, color='#2E2E2E'):
        """诊断 - 放大镜"""
        s = size
        lw = 2.5
        # 镜框
        canvas.create_oval(x+s*0.12, y+s*0.12, x+s*0.58, y+s*0.58, outline=color, width=lw)
        # 镜柄
        canvas.create_line(x+s*0.52, y+s*0.52, x+s*0.85, y+s*0.85, fill=color, width=lw, capstyle='round')
        # 高光
        canvas.create_arc(x+s*0.18, y+s*0.18, x+s*0.4, y+s*0.4, start=120, extent=60, style='arc', outline=color, width=1.5)
    
    @staticmethod
    def draw_edit(canvas, x, y, size=48, color='#2E2E2E'):
        """标点修复 - 铅笔"""
        s = size
        lw = 2.5
        # 笔身
        canvas.create_line(x+s*0.2, y+s*0.8, x+s*0.7, y+s*0.3, fill=color, width=lw, capstyle='round')
        # 笔尖
        canvas.create_polygon(
            x+s*0.15, y+s*0.85,
            x+s*0.2, y+s*0.8,
            x+s*0.25, y+s*0.85,
            fill=color, outline=''
        )
        # 笔头
        canvas.create_line(x+s*0.7, y+s*0.3, x+s*0.8, y+s*0.2, fill=color, width=lw, capstyle='round')
        canvas.create_line(x+s*0.75, y+s*0.35, x+s*0.85, y+s*0.25, fill=color, width=lw, capstyle='round')
    
    @staticmethod
    def draw_file(canvas, x, y, size=48, color='#2E2E2E'):
        """文件图标"""
        s = size
        lw = 2
        # 文件主体
        points = [
            x+s*0.2, y+s*0.1,   # 左上
            x+s*0.2, y+s*0.9,   # 左下
            x+s*0.8, y+s*0.9,   # 右下
            x+s*0.8, y+s*0.3,   # 右上（折角下）
            x+s*0.6, y+s*0.1,   # 折角
        ]
        canvas.create_polygon(points, fill='', outline=color, width=lw)
        # 折角线
        canvas.create_line(x+s*0.6, y+s*0.1, x+s*0.6, y+s*0.3, fill=color, width=lw)
        canvas.create_line(x+s*0.6, y+s*0.3, x+s*0.8, y+s*0.3, fill=color, width=lw)
    
    @staticmethod
    def draw_check(canvas, x, y, size=32, color='#7CB87C'):
        """勾选"""
        s = size
        canvas.create_line(x+s*0.15, y+s*0.5, x+s*0.4, y+s*0.75, fill=color, width=3, capstyle='round')
        canvas.create_line(x+s*0.4, y+s*0.75, x+s*0.85, y+s*0.25, fill=color, width=3, capstyle='round')


class FileInputField(tk.Frame):
    """文件输入框 - 带明显容器"""
    
    def __init__(self, parent, label_text, placeholder, variable, command, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.variable = variable
        self.command = command
        self.placeholder = placeholder
        
        # 标签
        tk.Label(
            self,
            text=label_text,
            font=get_font(11),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY,
            width=4,
            anchor='w'
        ).pack(side='left')
        
        # 输入框容器
        self.container = tk.Frame(
            self,
            bg=Theme.INPUT_BG,
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.PRIMARY,
            highlightthickness=1
        )
        self.container.pack(side='left', fill='x', expand=True, padx=(Theme.SPACE_SM, 0))
        
        inner = tk.Frame(self.container, bg=Theme.INPUT_BG)
        inner.pack(fill='both', expand=True, padx=Theme.SPACE_MD, pady=Theme.SPACE_SM + 2)
        
        # 文件名显示
        self.filename_label = tk.Label(
            inner,
            text="未选择",
            font=get_font(11),
            bg=Theme.INPUT_BG,
            fg=Theme.TEXT_MUTED,
            anchor='w'
        )
        self.filename_label.pack(side='left', fill='x', expand=True)
        
        # 分隔线
        tk.Frame(inner, bg=Theme.BORDER, width=1).pack(side='left', fill='y', padx=Theme.SPACE_MD)
        
        # 操作按钮
        self.action_btn = tk.Label(
            inner,
            text=placeholder,
            font=get_font(10),
            bg=Theme.INPUT_BG,
            fg=Theme.PRIMARY,
            cursor='hand2'
        )
        self.action_btn.pack(side='right')
        
        # 绑定点击
        for widget in [self.container, inner, self.filename_label, self.action_btn]:
            widget.bind('<Button-1>', self._on_click)
            widget.configure(cursor='hand2')
        
        # 悬停效果
        self.container.bind('<Enter>', lambda e: self.container.configure(highlightbackground='#D0CCC6'))
        self.container.bind('<Leave>', lambda e: self.container.configure(highlightbackground=Theme.BORDER))
        
        # 监听变量
        self.variable.trace_add('write', self._update_display)
    
    def _on_click(self, event=None):
        if self.command:
            self.command()
    
    def _update_display(self, *args):
        path = self.variable.get()
        if path:
            # 显示文件名，路径过长则截断
            filename = Path(path).name
            if len(filename) > 40:
                filename = filename[:37] + "..."
            self.filename_label.configure(text=filename, fg=Theme.TEXT)
        else:
            self.filename_label.configure(text="未选择", fg=Theme.TEXT_MUTED)


class SelectableCard(tk.Frame):
    """可选择的卡片 - 大图标版"""
    
    def __init__(self, parent, title, description, value, variable,
                 icon_draw_func=None, is_featured=False, command=None, **kwargs):
        
        bg_color = Theme.CARD_ALT if is_featured else Theme.CARD
        super().__init__(parent, bg=bg_color, **kwargs)
        
        self.value = value
        self.variable = variable
        self.command = command
        self.is_featured = is_featured
        self.bg_color = bg_color
        self.selected = False
        
        # 边框
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        # 内容 - 水平布局：左图标 + 右文字
        content = tk.Frame(self, bg=bg_color)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
        
        # 左侧：图标
        if icon_draw_func:
            icon_size = 56 if is_featured else 48
            self.icon_canvas = tk.Canvas(
                content,
                width=icon_size + 8,
                height=icon_size + 8,
                bg=bg_color,
                highlightthickness=0
            )
            self.icon_canvas.pack(side='left', padx=(0, Theme.SPACE_MD))
            icon_draw_func(self.icon_canvas, 4, 4, icon_size, Theme.TEXT)
            self._bind_click(self.icon_canvas)
        
        # 右侧：文字区域
        text_frame = tk.Frame(content, bg=bg_color)
        text_frame.pack(side='left', fill='both', expand=True)
        
        # 标题行（标题 + 推荐标签）
        title_row = tk.Frame(text_frame, bg=bg_color)
        title_row.pack(fill='x', anchor='w')
        
        title_size = 16 if is_featured else 14
        self.title_label = tk.Label(
            title_row,
            text=title,
            font=get_font(title_size, 'bold'),
            bg=bg_color,
            fg=Theme.TEXT,
            anchor='w'
        )
        self.title_label.pack(side='left')
        
        # 推荐标签
        if is_featured:
            tag = tk.Label(
                title_row,
                text=" 推荐 ",
                font=get_font(10, 'bold'),
                bg=Theme.PRIMARY,
                fg='white',
                padx=10,
                pady=3
            )
            tag.pack(side='left', padx=(Theme.SPACE_SM, 0))
            self._bind_click(tag)
        
        self._bind_click(title_row)
        
        # 描述
        desc_size = 12 if is_featured else 11
        self.desc_label = tk.Label(
            text_frame,
            text=description,
            font=get_font(desc_size),
            bg=bg_color,
            fg=Theme.TEXT_SECONDARY,
            anchor='w',
            justify='left'
        )
        self.desc_label.pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        # 绑定事件
        self._bind_click(self)
        self._bind_click(content)
        self._bind_click(text_frame)
        self._bind_click(self.title_label)
        self._bind_click(self.desc_label)
        
        # 监听变量
        self.variable.trace_add('write', self._on_variable_change)
        self._update_style()
    
    def _bind_click(self, widget):
        widget.bind('<Button-1>', self._on_click)
        widget.bind('<Enter>', self._on_enter)
        widget.bind('<Leave>', self._on_leave)
        widget.configure(cursor='hand2')
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
        if self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _on_variable_change(self, *args):
        self._update_style()
    
    def _update_style(self):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(highlightbackground=Theme.BORDER_SELECTED, highlightthickness=2)
        else:
            self.configure(highlightbackground=Theme.BORDER, highlightthickness=1)


class PresetCard(tk.Frame):
    """格式预设卡片"""
    
    def __init__(self, parent, text, value, variable, **kwargs):
        super().__init__(parent, bg=Theme.CARD, **kwargs)
        
        self.value = value
        self.variable = variable
        self.selected = False
        
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        self.label = tk.Label(
            self,
            text=text,
            font=get_font(12),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD
        )
        self.label.pack()
        
        # 绑定
        for widget in [self, self.label]:
            widget.bind('<Button-1>', self._on_click)
            widget.bind('<Enter>', self._on_enter)
            widget.bind('<Leave>', self._on_leave)
            widget.configure(cursor='hand2')
        
        self.variable.trace_add('write', self._update_style)
        self._update_style()
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _update_style(self, *args):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(bg=Theme.PRIMARY_LIGHT, highlightbackground=Theme.PRIMARY, highlightthickness=2)
            self.label.configure(bg=Theme.PRIMARY_LIGHT, fg=Theme.TEXT, font=get_font(12, 'bold'))
        else:
            self.configure(bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
            self.label.configure(bg=Theme.CARD, fg=Theme.TEXT, font=get_font(12))
    
    def set_enabled(self, enabled):
        if enabled:
            self.label.configure(fg=Theme.TEXT, cursor='hand2')
            self.configure(cursor='hand2')
        else:
            self.label.configure(fg=Theme.TEXT_MUTED, cursor='arrow')
            self.configure(cursor='arrow', highlightbackground=Theme.BORDER_LIGHT)


class CollapsibleLog(tk.Frame):
    """可折叠的日志区域"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.expanded = False
        
        # 折叠条
        self.toggle_bar = tk.Frame(self, bg='#E8E4DE', height=36)
        self.toggle_bar.pack(fill='x')
        self.toggle_bar.pack_propagate(False)
        
        self.toggle_btn = tk.Label(
            self.toggle_bar,
            text="＋  展开运行日志",
            font=get_font(11),
            bg='#E8E4DE',
            fg=Theme.TEXT_SECONDARY,
            cursor='hand2'
        )
        self.toggle_btn.pack(side='left', padx=Theme.SPACE_MD, pady=Theme.SPACE_SM)
        self.toggle_btn.bind('<Button-1>', self._toggle)
        self.toggle_bar.bind('<Button-1>', self._toggle)
        self.toggle_bar.configure(cursor='hand2')
        
        # 日志面板
        self.log_panel = tk.Frame(self, bg=Theme.LOG_BG)
        
        # 日志文本
        self.log_text = tk.Text(
            self.log_panel,
            font=('Consolas', 11),
            bg=Theme.LOG_BG,
            fg=Theme.LOG_TEXT,
            relief='flat',
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD,
            wrap='word',
            height=10,
            highlightthickness=0,
            insertbackground=Theme.LOG_TEXT
        )
        self.log_text.pack(side='left', fill='both', expand=True)
        
        # 配置颜色标签
        self.log_text.tag_configure('info', foreground=Theme.LOG_TEXT)
        self.log_text.tag_configure('success', foreground=Theme.LOG_SUCCESS)
        self.log_text.tag_configure('warning', foreground=Theme.LOG_WARNING)
        self.log_text.tag_configure('error', foreground=Theme.LOG_ERROR)
    
    def _toggle(self, event=None):
        self.expanded = not self.expanded
        if self.expanded:
            self.log_panel.pack(fill='both', expand=True)
            self.toggle_btn.configure(text="－  收起运行日志")
        else:
            self.log_panel.pack_forget()
            self.toggle_btn.configure(text="＋  展开运行日志")
    
    def log(self, message, tag='info'):
        self.log_text.insert(tk.END, message + "\n", tag)
        self.log_text.see(tk.END)
    
    def clear(self):
        self.log_text.delete(1.0, tk.END)


class ResultPanel(tk.Frame):
    """结果反馈面板"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        # 占位状态
        self.placeholder = tk.Label(
            self,
            text="处理结果将在此处显示",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            pady=Theme.SPACE_XL
        )
        self.placeholder.pack()
        
        # 结果卡片
        self.result_card = tk.Frame(self, bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
        self.result_content = tk.Frame(self.result_card, bg=Theme.CARD)
        self.result_content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
    
    def show_success(self, message, filepath=None):
        self.placeholder.pack_forget()
        
        for widget in self.result_content.winfo_children():
            widget.destroy()
        
        # 成功图标 + 消息
        header = tk.Frame(self.result_content, bg=Theme.CARD)
        header.pack(fill='x', anchor='w')
        
        icon_canvas = tk.Canvas(header, width=36, height=36, bg=Theme.CARD, highlightthickness=0)
        icon_canvas.pack(side='left')
        Icons.draw_check(icon_canvas, 2, 2, 32, Theme.LOG_SUCCESS)
        
        tk.Label(
            header,
            text=message,
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(side='left', padx=(Theme.SPACE_SM, 0))
        
        if filepath:
            tk.Label(
                self.result_content,
                text=f"输出文件：{filepath}",
                font=get_font(11),
                bg=Theme.CARD,
                fg=Theme.TEXT_SECONDARY,
                anchor='w'
            ).pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))
    
    def show_diagnosis(self, results):
        self.placeholder.pack_forget()
        
        for widget in self.result_content.winfo_children():
            widget.destroy()
        
        tk.Label(
            self.result_content,
            text="诊断报告",
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(fill='x', anchor='w', pady=(0, Theme.SPACE_MD))
        
        total = 0
        categories = [
            ('标点问题', results.get('punctuation', [])),
            ('序号问题', results.get('numbering', [])),
            ('段落问题', results.get('paragraph', [])),
            ('字体问题', results.get('font', [])),
        ]
        
        for name, issues in categories:
            count = len(issues)
            total += count
            
            row = tk.Frame(self.result_content, bg=Theme.CARD)
            row.pack(fill='x', pady=3)
            
            tk.Label(
                row,
                text=name,
                font=get_font(12),
                bg=Theme.CARD,
                fg=Theme.TEXT,
                width=10,
                anchor='w'
            ).pack(side='left')
            
            count_color = Theme.LOG_WARNING if count > 0 else Theme.LOG_SUCCESS
            tk.Label(
                row,
                text=f"{count} 处" if count > 0 else "无问题",
                font=get_font(12),
                bg=Theme.CARD,
                fg=count_color,
                anchor='w'
            ).pack(side='left')
        
        tk.Frame(self.result_content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_MD)
        
        summary_color = Theme.LOG_SUCCESS if total == 0 else Theme.LOG_WARNING
        summary_text = "文档格式规范，未发现问题" if total == 0 else f"共发现 {total} 处格式问题"
        
        tk.Label(
            self.result_content,
            text=summary_text,
            font=get_font(13, 'bold'),
            bg=Theme.CARD,
            fg=summary_color,
            anchor='w'
        ).pack(fill='x', anchor='w')
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))
    
    def reset(self):
        self.result_card.pack_forget()
        for widget in self.result_content.winfo_children():
            widget.destroy()
        self.placeholder.pack()


class DocFormatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("公文格式处理工具")
        self.root.geometry("750x900")
        self.root.minsize(680, 750)
        self.root.configure(bg=Theme.BG)
        
        # 变量
        self.input_file = tk.StringVar()
        self.output_file = tk.StringVar()
        self.operation = tk.StringVar(value="smart")
        self.preset = tk.StringVar(value="official")
        
        self.preset_cards = []
        
        self.create_widgets()
    
    def create_widgets(self):
        """构建界面"""
        # 主容器 - 带滚动
        container = tk.Frame(self.root, bg=Theme.BG)
        container.pack(fill='both', expand=True)
        
        # Canvas + 自定义滚动条
        self.canvas = tk.Canvas(container, bg=Theme.BG, highlightthickness=0)
        self.scrollbar_canvas = tk.Canvas(container, width=14, bg=Theme.BG, highlightthickness=0)
        
        self.canvas.pack(side='left', fill='both', expand=True)
        self.scrollbar_canvas.pack(side='right', fill='y')
        
        # 内容Frame
        self.main_frame = tk.Frame(self.canvas, bg=Theme.BG)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.main_frame, anchor='nw')
        
        # 绑定滚动
        self.main_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.root.bind_all('<MouseWheel>', self._on_mousewheel)
        self.scrollbar_canvas.bind('<Button-1>', self._on_scrollbar_click)
        self.scrollbar_canvas.bind('<B1-Motion>', self._on_scrollbar_drag)
        
        # 内容区域
        content = tk.Frame(self.main_frame, bg=Theme.BG)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_XL, pady=Theme.SPACE_LG)
        
        # ===== 1. 头部 =====
        tk.Label(
            content,
            text="公文格式处理工具",
            font=get_font(24, 'bold'),
            bg=Theme.BG,
            fg=Theme.TEXT
        ).pack(anchor='w', pady=(0, Theme.SPACE_XL))
        
        # ===== 2. 文件选择区 =====
        file_section = tk.Frame(content, bg=Theme.BG)
        file_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        self.input_field = FileInputField(
            file_section,
            label_text="输入",
            placeholder="点击选择需要修改的文档",
            variable=self.input_file,
            command=self.browse_input
        )
        self.input_field.pack(fill='x', pady=(0, Theme.SPACE_SM))
        
        self.output_field = FileInputField(
            file_section,
            label_text="输出",
            placeholder="文档修改后的储存位置",
            variable=self.output_file,
            command=self.browse_output
        )
        self.output_field.pack(fill='x')
        
        # 分隔
        tk.Frame(content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_LG)
        
        # ===== 3. 功能选择区 =====
        mode_section = tk.Frame(content, bg=Theme.BG)
        mode_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # 大卡片 - 智能一键处理
        smart_card = SelectableCard(
            mode_section,
            title="智能一键处理",
            description="自动修复标点符号，并应用标准格式规范，一步到位完成文档处理",
            value="smart",
            variable=self.operation,
            icon_draw_func=Icons.draw_magic,
            is_featured=True,
            command=self._on_mode_change
        )
        smart_card.pack(fill='x', pady=(0, Theme.SPACE_MD))
        
        # 两个小卡片
        small_cards = tk.Frame(mode_section, bg=Theme.BG)
        small_cards.pack(fill='x')
        small_cards.columnconfigure(0, weight=1)
        small_cards.columnconfigure(1, weight=1)
        
        diag_card = SelectableCard(
            small_cards,
            title="格式诊断",
            description="仅分析文档问题，不修改文件",
            value="analyze",
            variable=self.operation,
            icon_draw_func=Icons.draw_search,
            command=self._on_mode_change
        )
        diag_card.grid(row=0, column=0, sticky='nsew', padx=(0, Theme.SPACE_SM))
        
        punct_card = SelectableCard(
            small_cards,
            title="标点修复",
            description="仅修复中英文标点混用",
            value="punctuation",
            variable=self.operation,
            icon_draw_func=Icons.draw_edit,
            command=self._on_mode_change
        )
        punct_card.grid(row=0, column=1, sticky='nsew')
        
        # ===== 4. 格式预设 =====
        preset_section = tk.Frame(content, bg=Theme.BG)
        preset_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        tk.Label(
            preset_section,
            text="格式预设",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY
        ).pack(anchor='w', pady=(0, Theme.SPACE_SM))
        
        preset_row = tk.Frame(preset_section, bg=Theme.BG)
        preset_row.pack(fill='x')
        
        presets = [
            ('official', 'GB/T 公文标准'),
            ('academic', '学术论文'),
            ('legal', '法律文书'),
        ]
        
        for i, (value, text) in enumerate(presets):
            card = PresetCard(preset_row, text, value, self.preset)
            card.pack(side='left', padx=(0 if i == 0 else Theme.SPACE_SM, 0))
            self.preset_cards.append(card)
        
        # ===== 5. 执行按钮 =====
        self.run_btn = tk.Frame(content, bg=Theme.PRIMARY, cursor='hand2')
        self.run_btn.pack(fill='x', pady=Theme.SPACE_LG)
        
        self.run_label = tk.Label(
            self.run_btn,
            text="开始处理",
            font=get_font(15, 'bold'),
            bg=Theme.PRIMARY,
            fg='white',
            pady=Theme.SPACE_MD + 2
        )
        self.run_label.pack()
        
        for widget in [self.run_btn, self.run_label]:
            widget.bind('<Button-1>', lambda e: self.run_operation())
            widget.bind('<Enter>', lambda e: self._btn_hover(True))
            widget.bind('<Leave>', lambda e: self._btn_hover(False))
        self.run_label.configure(cursor='hand2')
        
        # ===== 6. 结果反馈区 =====
        self.result_panel = ResultPanel(content)
        self.result_panel.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # ===== 7. 日志区 =====
        self.log_panel = CollapsibleLog(content)
        self.log_panel.pack(fill='x', pady=(Theme.SPACE_MD, 0))
        
        # 初始化
        self._on_mode_change()
        self.log_panel.log("工具已就绪，请选择文件", 'info')
    
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
        self._draw_scrollbar()
    
    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)
        self._draw_scrollbar()
    
    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')
        self._draw_scrollbar()
    
    def _draw_scrollbar(self):
        """绘制自定义滚动条"""
        self.scrollbar_canvas.delete('all')
        
        try:
            top, bottom = self.canvas.yview()
        except:
            return
        
        if bottom - top >= 0.99:
            return
        
        w = 14
        h = self.scrollbar_canvas.winfo_height()
        
        if h < 10:
            return
        
        bar_h = max(40, (bottom - top) * h)
        bar_y = top * (h - bar_h)
        
        # 轨道
        self.scrollbar_canvas.create_rectangle(
            4, 8, w - 4, h - 8,
            fill='#E8E4DE', outline=''
        )
        
        # 滑块（更深的颜色）
        self.scrollbar_canvas.create_rectangle(
            4, bar_y + 8, w - 4, bar_y + bar_h - 8,
            fill='#A09890', outline=''
        )
    
    def _on_scrollbar_click(self, event):
        """滚动条点击"""
        try:
            h = self.scrollbar_canvas.winfo_height()
            fraction = event.y / h
            self.canvas.yview_moveto(fraction)
            self._draw_scrollbar()
        except:
            pass
    
    def _on_scrollbar_drag(self, event):
        """滚动条拖动"""
        self._on_scrollbar_click(event)
    
    def _btn_hover(self, is_hover):
        color = Theme.PRIMARY_HOVER if is_hover else Theme.PRIMARY
        self.run_btn.configure(bg=color)
        self.run_label.configure(bg=color)
    
    def _on_mode_change(self):
        mode = self.operation.get()
        enabled = mode in ('smart',)
        for card in self.preset_cards:
            card.set_enabled(enabled)
    
    def browse_input(self):
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.input_file.set(filename)
            p = Path(filename)
            output_name = f"{p.stem}_processed{p.suffix}"
            self.output_file.set(str(p.parent / output_name))
            self.log_panel.log(f"已选择: {p.name}", 'info')
            self.result_panel.reset()
    
    def browse_output(self):
        filename = filedialog.asksaveasfilename(
            title="保存为",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx")]
        )
        if filename:
            self.output_file.set(filename)
    
    def run_operation(self):
        input_path = self.input_file.get().strip()
        output_path = self.output_file.get().strip()
        mode = self.operation.get()
        
        if not input_path:
            messagebox.showerror("提示", "请先选择输入文件")
            return
        
        if not os.path.exists(input_path):
            messagebox.showerror("错误", "文件不存在")
            return
        
        if mode != 'analyze' and not output_path:
            messagebox.showerror("提示", "请指定输出文件")
            return
        
        self.run_btn.configure(bg=Theme.TEXT_MUTED)
        self.run_label.configure(bg=Theme.TEXT_MUTED, text="处理中...")
        
        thread = threading.Thread(
            target=self._do_operation,
            args=(input_path, output_path, mode)
        )
        thread.start()
    
    def _do_operation(self, input_path, output_path, mode):
        try:
            from docx import Document
            
            self.log_panel.log(f"\n{'─' * 35}", 'info')
            self.log_panel.log(f"开始处理: {Path(input_path).name}", 'info')
            
            if mode == 'analyze':
                doc = Document(input_path)
                results = {
                    'punctuation': analyze_punctuation(doc),
                    'numbering': analyze_numbering(doc),
                    'paragraph': analyze_paragraph_format(doc),
                    'font': analyze_font(doc)
                }
                self.root.after(0, lambda: self.result_panel.show_diagnosis(results))
                self.log_panel.log("诊断完成", 'success')
                
            elif mode == 'punctuation':
                self._run_punctuation(input_path, output_path)
                self.root.after(0, lambda: self.result_panel.show_success(
                    "标点修复完成", Path(output_path).name
                ))
                
            elif mode == 'smart':
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    temp_path = tmp.name
                
                self.log_panel.log("步骤 1/2: 修复标点...", 'info')
                self._run_punctuation(input_path, temp_path, quiet=True)
                
                self.log_panel.log("步骤 2/2: 应用格式...", 'info')
                self._run_format(temp_path, output_path)
                
                os.unlink(temp_path)
                
                self.root.after(0, lambda: self.result_panel.show_success(
                    "处理完成", Path(output_path).name
                ))
            
            self.log_panel.log("全部完成", 'success')
            
            if mode != 'analyze':
                self.root.after(0, lambda: messagebox.showinfo(
                    "完成", f"文件已保存至:\n{output_path}"
                ))
        
        except Exception as e:
            self.log_panel.log(f"错误: {str(e)}", 'error')
            import traceback
            self.log_panel.log(traceback.format_exc(), 'error')
            self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        
        finally:
            self.root.after(0, self._reset_btn)
    
    def _reset_btn(self):
        self.run_btn.configure(bg=Theme.PRIMARY)
        self.run_label.configure(bg=Theme.PRIMARY, text="开始处理")
    
    def _run_punctuation(self, input_path, output_path, quiet=False):
        from docx import Document
        from scripts.punctuation import process_paragraph
        
        doc = Document(input_path)
        changes = 0
        
        for para in doc.paragraphs:
            if process_paragraph(para):
                changes += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if process_paragraph(para):
                            changes += 1
        
        doc.save(output_path)
        if not quiet:
            self.log_panel.log(f"修复了 {changes} 处标点", 'success')
    
    def _run_format(self, input_path, output_path):
        preset_name = self.preset.get()
        
        import io
        old_stdout = sys.stdout
        sys.stdout = io.StringIO()
        
        try:
            format_document(input_path, output_path, preset_name)
        finally:
            sys.stdout = old_stdout
        
        self.log_panel.log(f"应用格式: {PRESETS[preset_name]['name']}", 'success')


def main():
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    root = tk.Tk()
    app = DocFormatApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
